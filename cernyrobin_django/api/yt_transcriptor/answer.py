import os

try:
    import openai
    import dotenv
    import requests
    from pytube import YouTube
    from docx import Document
except ImportError as e:
  print("Missing dependencies detected. Run pip install -r requirements.txt to install...")
  print(e)

MODEL = "gpt-3.5-turbo-16k-0613" # The best model for this application - large text, but stay cheap

try:
    dotenv.load_dotenv()
except Exception as e:
    print("Failed to load dotenv file. This should only happen in a docker container!")
    print(e)

STUCKINVIM_KEY = os.getenv("STUCKINVIM_KEY")
print("STUCKINVIM_KEY: " + STUCKINVIM_KEY)
TEMPERATURE = 0.35 # Tweaked manually
REGEN_TEMPERATURE = 0.45
MAX_TOKENS = 1250

if not STUCKINVIM_KEY:
    print("Missing API key. Please set the STUCKINVIM_KEY environment variable.")
    exit(1)

print("")

API_KEY = None

try:
    result = requests.get("http://getkey.stuckinvim.com/api/data?api_key=" + STUCKINVIM_KEY)

    if result is not None:
        result = result.json()
    else:
        "Failed to get API key from StuckInVim API."

    if result.get("status", "400") == "200":
        API_KEY = result["key"]
    else:
        raise Exception("Failed to get API key from StuckInVim API.")

    openai.api_key = API_KEY
except Exception as e:
    print("Failed to get OpenAI API key. Script will continue to run with limited features")

def get_video_description(video_url):
    if video_url.startswith("C:") or video_url.startswith("\\") or video_url.startswith("/"):
        return None
    elif not video_url.startswith("http") and not video_url.startswith("www") and not video_url.startswith("youtube"):
        video_url = "https://www.youtube.com/watch?v=" + video_url
        
    try:
        yt = YouTube(video_url)
        return yt.initial_data["engagementPanels"][1]["engagementPanelSectionListRenderer"]["content"]["structuredDescriptionContentRenderer"]["items"][1]["expandableVideoDescriptionBodyRenderer"]["attributedDescriptionBodyText"]["content"]
    except Exception as e:
        print(e)
        return None

def chatbot(questions_path, transcript_path, save_path, summary_save_path, youtube_url=None, is_regen=False):
    temperature = 0

    if is_regen:
        temperature = REGEN_TEMPERATURE
    else:
        temperature = TEMPERATURE

    questions = ""
    if questions_path:
        try:
            doc = Document(questions_path)
        except:
            print("Error")
        if questions_path.endswith(".docx") or questions_path.endswith(".doc"):
            for paragraph in doc.paragraphs:
                questions += paragraph.text
        elif questions_path.endswith(".txt"):
            with open(questions_path, encoding="utf-8") as txt:
                questions = txt.read()
    elif youtube_url:
        print("Getting from URL")
        questions = get_video_description(youtube_url)
    else:
        print("No questions provided.")
        exit(1)

    if questions is None or questions == "":
        print("Failed to get questions from description.")
        exit(1) 

    transcript = ""
    with open(transcript_path, encoding="utf-8") as txt:
       transcript = txt.read()

    system_message_path = os.path.join(os.path.dirname(__file__), "system_message.txt")
    system_message = ""
    summary_prompt_path = os.path.join(os.path.dirname(__file__), "summary_prompt.txt")
    summary_prompt = ""

    if os.path.exists(system_message_path):
        with open(system_message_path, encoding="utf-8") as txt:    
            system_message = txt.read()

    if os.path.exists(summary_prompt_path):
        with open(summary_prompt_path, encoding="utf-8") as txt:
            summary_prompt = txt.read()

    messages = [
        {"role": "system", "content": system_message},
    ]

    messages.append({"role": "user", "content": transcript})

    print("Sending prompt with transcript...")

    response = openai.chat.completions.create(
        model=MODEL,
        messages=messages,
        temperature=temperature,
        max_tokens=MAX_TOKENS
    )

    print("Sending prompt with questions...")

    messages.append({"role": "user", "content": questions})
 
    response = openai.chat.completions.create(
        model=MODEL,
        messages=messages,
        temperature=temperature,
        max_tokens=MAX_TOKENS
    )

    try:
        response_text = response.choices[0].message.content

        if save_path.endswith(".docx") or save_path.endswith(".doc"):
            doc = Document()
            doc.add_paragraph(response_text)
            doc.save(save_path)
        elif save_path.endswith(".txt"):
            with open(save_path, "w", encoding="utf-8") as txt:
                txt.write(response_text)
    except:
        print("Error while saving answers. Will print to console instead.")

        print(response)

    print("Sending summary prompt...")
    messages.append({"role": "user", "content": summary_prompt})

    response = openai.chat.completions.create(
        model=MODEL,
        messages=messages,
        temperature=temperature,
        max_tokens=MAX_TOKENS
    )

    # summary_save_path = os.path.join(os.path.dirname(save_path), os.path.basename(save_path).split(".")[0] + "_summary.txt")

    print("Saving summary to " + summary_save_path)

    try:
        response_text = response.choices[0].message.content

        if summary_save_path.endswith(".docx") or summary_save_path.endswith(".doc"):
            doc = Document()
            doc.add_paragraph(response_text)
            doc.save(summary_save_path)
        elif summary_save_path.endswith(".txt"):
            with open(summary_save_path, "w", encoding="utf-8") as txt:
                txt.write(response_text)
    except:
        print("Error while saving summary. Will print to console instead.")

        print(response)